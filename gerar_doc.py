"""
Gera o documento de entrega da atividade de revisão de arquitetura.
Execute com: venv/Scripts/python.exe gerar_doc.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Margens ───────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(3)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2)

# ── Estilos auxiliares ────────────────────────────────────────
def set_font(run, bold=False, size=12, color=None, italic=False):
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Arial"
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    p.style.font.name = "Arial"
    for run in p.runs:
        run.font.name = "Arial"
        if level == 1:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)
        elif level == 2:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x1d, 0x4e, 0x89)
        else:
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def body(text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_font(run, bold=bold, italic=italic, size=size)
    return p

def bullet(text, size=11):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_font(run, size=size)
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    # fundo cinza claro via sombreamento
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F1F5F9")
    pPr.append(shd)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(10)
        # cabeçalho azul escuro
        tc_pr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "1e40af")
        tc_pr.append(shd)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xff, 0xff, 0xff)

    for row_data in rows:
        row = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            row[i].text = cell_text
            for p in row[i].paragraphs:
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(10)
    return table

# ══════════════════════════════════════════════════════════════
# CAPA
# ══════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("AgroVision AI")
set_font(run, bold=True, size=24, color=(0x1e, 0x40, 0xaf))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Revisão de Arquitetura, Segurança e Implementação de Scraping")
set_font(run, bold=False, size=14, color=(0x33, 0x33, 0x33))

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Sistema de Monitoramento Rural com Detecção por IA")
set_font(run, italic=True, size=12, color=(0x55, 0x55, 0x55))

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"Data de entrega: {datetime.date.today().strftime('%d/%m/%Y')}")
set_font(run, size=11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# INTRODUÇÃO
# ══════════════════════════════════════════════════════════════
heading("Introdução")
body(
    "O projeto AgroVision AI é um sistema de monitoramento rural que utiliza visão computacional "
    "(YOLO) para detectar pessoas, veículos e outros objetos em tempo real via câmera. O sistema "
    "conta com uma interface web, um backend em FastAPI, banco de dados SQLite, integração com o "
    "modelo de linguagem Llama via Ollama e uma camada de scraping de dados públicos. Este documento "
    "apresenta a revisão crítica da arquitetura, os riscos de segurança identificados, as melhorias "
    "realizadas sobre código gerado com IA e a implementação da camada de web scraping."
)

# ══════════════════════════════════════════════════════════════
# PARTE 1
# ══════════════════════════════════════════════════════════════
doc.add_paragraph()
heading("Parte 1 — Revisão da Arquitetura")

heading("1.1 Estrutura de arquivos do projeto", level=2)
code_block(
    "agrovision_ia/\n"
    "├── app.py            ← monolito: rotas, lógica, banco, IA e integrações\n"
    "├── scraper.py        ← camada de scraping (isolada)\n"
    "├── templates/        ← frontend (HTML + Jinja2)\n"
    "├── static/           ← frontend (CSS + JavaScript)\n"
    "├── dataset_agro/     ← dados para treino do YOLO\n"
    "├── detections.db     ← banco de dados SQLite\n"
    "└── yolo11n.pt        ← pesos do modelo YOLO"
)

heading("1.2 Diagnóstico por camada", level=2)
add_table(
    ["Camada", "Arquivo / Local", "Situação", "Observação"],
    [
        ["Frontend", "index.html + dashboard.js", "Aceitável",
         "Responsável apenas por exibir dados e fazer chamadas à API. Sem regra de negócio indevida."],
        ["Backend / API", "app.py (arquivo único)", "Ruim",
         "Rotas, lógica de negócio, acesso a banco, IA e integrações externas convivem no mesmo arquivo, sem separação de módulos."],
        ["Banco de Dados", "app.py — linhas 226 a 303", "Ruim",
         "Funções de acesso ao banco definidas no mesmo arquivo das rotas. Nova conexão SQLite aberta e fechada a cada operação, sem pool."],
        ["Serviços Internos", "app.py — process_stream(), warmup_ollama()", "Ruim",
         "Serviços rodando em threads separadas, mas definidos no mesmo módulo principal, dificultando manutenção."],
        ["Camada de IA / YOLO", "app.py — process_stream()", "Ruim",
         "A chamada ao modelo (model(frame, ...)) está misturada com leitura de câmera, desenho de bounding boxes e escrita no banco."],
        ["Integração Externa", "app.py — Ollama, Windy, yt-dlp", "Regular",
         "Cada integração possui funções dedicadas, mas todas residem no mesmo arquivo de rotas."],
        ["Web Scraping", "scraper.py (módulo próprio)", "Bom",
         "Implementado como módulo isolado. O app.py apenas importa e expõe os dados via endpoints dedicados /scraping/*."],
    ]
)

doc.add_paragraph()
heading("1.3 Respostas às perguntas da atividade", level=2)

heading("A interface está apenas exibindo dados ou também possui regra de negócio indevida?", level=3)
body(
    "A interface exibe dados e realiza chamadas à API de forma adequada. O único ponto de atenção é "
    "a detecção do tipo de URL (HLS, YouTube, RTSP) feita no JavaScript no lado do cliente, lógica "
    "que poderia estar centralizada no backend. No geral, o frontend não carrega responsabilidades indevidas."
)

heading("O backend concentra a lógica principal do sistema?", level=3)
body(
    "Sim, toda a lógica está no backend — porém sem organização. O arquivo app.py acumula "
    "responsabilidades que deveriam estar em módulos separados: repositório de banco, serviço de "
    "câmera, serviço de IA, integrações externas e rotas HTTP. Isso dificulta testes, manutenção e "
    "expansão do sistema."
)

heading("O acesso ao banco está isolado em uma camada própria?", level=3)
body(
    "Não. As funções init_db, save_event, list_events e get_last_event estão definidas no mesmo "
    "arquivo das rotas HTTP. Além disso, save_event é chamada diretamente dentro de process_stream, "
    "uma thread de processamento de vídeo em tempo real, misturando I/O de banco com processamento "
    "de câmera. O ideal seria uma camada de repositório separada."
)

heading("A chamada ao modelo de IA/YOLO está separada da regra de negócio?", level=3)
body(
    "Não. A linha model(frame, conf=CONFIDENCE_THRESHOLD, verbose=False) está dentro de "
    "process_stream, junto com: leitura de frame da câmera, desenho de bounding boxes, lógica de "
    "cooldown de alertas e escrita no banco. Uma boa separação exigiria um serviço dedicado de "
    "inferência que recebesse o frame e devolvesse as detecções."
)

heading("A nova camada de scraping foi implementada como serviço separado?", level=3)
body(
    "Sim. O arquivo scraper.py é um módulo independente, com funções puras (fetch_weather e "
    "fetch_agro_news), cache próprio, rate limiting e tratamento de erro isolado. O app.py importa "
    "o módulo e expõe os dados através de dois endpoints GET dedicados (/scraping/weather e "
    "/scraping/news), sem misturar a lógica de scraping com rotas ou controllers."
)

# ══════════════════════════════════════════════════════════════
# PARTE 2
# ══════════════════════════════════════════════════════════════
doc.add_page_break()
heading("Parte 2 — Revisão de Segurança")

heading("2.1 Riscos identificados", level=2)
add_table(
    ["#", "Risco", "Localização", "Severidade", "Descrição"],
    [
        ["1", "Rotas sem autenticação", "Todos os endpoints", "Alta",
         "Qualquer pessoa na rede pode trocar câmera, selecionar modelo Ollama, acessar o feed de vídeo e consultar eventos."],
        ["2", "SSRF (Server-Side Request Forgery)", "/camera/select e /stream/resolve", "Alta",
         "A URL enviada pelo usuário é repassada diretamente ao cv2.VideoCapture e ao yt-dlp sem validação de esquema ou domínio. Permite acesso a recursos internos da rede."],
        ["3", "Mensagens de erro técnicas expostas", "app.py linhas 484, 510, 714", "Média",
         "Exceções completas são retornadas ao cliente via str(exc), podendo revelar caminhos internos, versões de bibliotecas e dados de configuração."],
        ["4", "Sem rate limiting", "Endpoints /chat e /scraping/*", "Média",
         "Os endpoints de chat e scraping podem ser chamados ilimitadamente, permitindo abuso do Ollama e sobrecarga da fonte de scraping."],
        ["5", "Injeção de modelo Ollama", "/model/select", "Média",
         "MODEL_NAME = req.model aceita qualquer string sem validar se o modelo existe ou é permitido, podendo carregar modelos não autorizados."],
        ["6", "SQL Injection", "app.py linhas 244–290", "Baixa / Mitigada",
         "Não existe risco: todas as queries utilizam parâmetros posicionais (?), protegendo contra injeção de SQL."],
        ["7", "Secrets no código", "Variáveis de ambiente (.env)", "OK",
         "As chaves WINDY_API_KEY e OLLAMA_URL são carregadas via python-dotenv, não estão hardcoded no código."],
    ]
)

doc.add_paragraph()
heading("2.2 Risco com IA, scraping e fontes externas", level=2)
body(
    "O sistema processa URLs externas (câmeras, streams de vídeo, YouTube) e dados de scraping "
    "sem validação prévia. Os principais vetores de risco identificados são:"
)
bullet("SSRF: um atacante pode fornecer IPs de serviços internos (ex: http://192.168.x.x) como fonte de câmera, e o servidor fará a requisição.")
bullet("DoS por stream: URLs que apontam para fontes de vídeo de alta largura de banda podem consumir todos os recursos do servidor.")
bullet("Dados malformados: arquivos de vídeo corrompidos ou streams malformados podem explorar vulnerabilidades no OpenCV ou FFmpeg.")
bullet("Scraping de fontes não confiáveis: o conteúdo HTML raspado é exibido na tela sem sanitização, abrindo risco de XSS caso o site externo seja comprometido.")

body(
    "Mitigação recomendada: validar esquema de URL (apenas rtsp://, http://, https://), bloquear "
    "intervalos de IP privado antes de passar ao VideoCapture, e escapar o conteúdo HTML do scraping "
    "antes de renderizar no browser.",
    italic=True
)

# ══════════════════════════════════════════════════════════════
# PARTE 3
# ══════════════════════════════════════════════════════════════
doc.add_page_break()
heading("Parte 3 — Melhoria do Código Gerado com IA")

body(
    "Os três trechos abaixo foram criados ou fortemente auxiliados por IA durante o desenvolvimento "
    "do projeto e apresentavam problemas de qualidade que foram identificados e corrigidos durante "
    "esta revisão."
)

# Trecho 1
doc.add_paragraph()
heading("Trecho 1 — Resolução de URL do YouTube (resolve_stream_url)", level=2)

heading("O que o código fazia originalmente", level=3)
body("Chamava o yt-dlp de forma síncrona, sem qualquer timeout, para resolver URLs do YouTube:")
code_block(
    "ydl_opts = {\"format\": \"best[ext=mp4]/best\", \"quiet\": True}\n"
    "with yt_dlp.YoutubeDL(ydl_opts) as ydl:\n"
    "    info = ydl.extract_info(source, download=False)\n"
    "    direct_url = info.get(\"url\") or info[\"formats\"][-1][\"url\"]"
)

heading("Problema encontrado", level=3)
body(
    "Se o yt-dlp travasse (rede lenta, live encerrada, URL inválida), a thread principal de "
    "processamento de câmera ficava bloqueada indefinidamente. O usuário via a mensagem "
    "'Resolvendo stream do YouTube...' para sempre, sem feedback e sem possibilidade de cancelar."
)

heading("O que foi melhorado", level=3)
code_block(
    "# Executa yt-dlp em thread separada com timeout de 30 segundos\n"
    "with ThreadPoolExecutor(max_workers=1) as ex:\n"
    "    future = ex.submit(_yt_dlp_extract, source)\n"
    "    try:\n"
    "        direct_url = future.result(timeout=30)\n"
    "    except FuturesTimeout:\n"
    "        set_camera_status(\"error\", \"Timeout ao resolver YouTube URL (>30s).\")\n"
    "        return source"
)

heading("Por que a nova versão é melhor", level=3)
body(
    "A resolução ocorre em uma thread separada com timeout garantido via concurrent.futures. "
    "Se o yt-dlp travar, FuturesTimeout é capturado após 30 segundos e o erro é comunicado "
    "ao usuário através do endpoint /camera/status, que o frontend consulta por polling. "
    "O sistema continua funcionando normalmente."
)

# Trecho 2
doc.add_paragraph()
heading("Trecho 2 — Stream MJPEG sem frame disponível (generate_mjpeg)", level=2)

heading("O que o código fazia originalmente", level=3)
body("Quando a câmera estava indisponível ou em processo de troca, o gerador aguardava em loop:")
code_block(
    "if frame is None:\n"
    "    await asyncio.sleep(0.1)\n"
    "    continue  # nada é enviado ao cliente"
)

heading("Problema encontrado", level=3)
body(
    "O browser aguardava o primeiro frame do stream MJPEG indefinidamente, exibindo o ícone "
    "de imagem quebrada ou um símbolo de carregamento. O usuário não recebia nenhum feedback "
    "visual sobre o estado do sistema, podendo confundir a ausência de imagem com uma falha "
    "geral da aplicação."
)

heading("O que foi melhorado", level=3)
code_block(
    "if frame is None:\n"
    "    yield (\n"
    "        b\"--frame\\r\\nContent-Type: image/jpeg\\r\\n\\r\\n\"\n"
    "        + _placeholder_frame()   # imagem com texto \"Aguardando câmera...\"\n"
    "        + b\"\\r\\n\"\n"
    "    )\n"
    "    await asyncio.sleep(1.0)\n"
    "    continue"
)

heading("Por que a nova versão é melhor", level=3)
body(
    "O stream MJPEG nunca é interrompido. Enquanto não há frame disponível, o browser exibe "
    "a imagem de placeholder com a mensagem 'Aguardando câmera...', fornecendo feedback claro "
    "ao usuário. A taxa de 1 frame por segundo durante a espera evita consumo desnecessário "
    "de CPU e banda."
)

# Trecho 3
doc.add_paragraph()
heading("Trecho 3 — Conexão de stream por URL (applyStreamUrl no JavaScript)", level=2)

heading("O que o código fazia originalmente", level=3)
body("Qualquer tipo de URL (HLS, YouTube, RTSP, MJPEG) era enviada ao backend:")
code_block(
    "// Toda URL ia para o backend, independente do tipo\n"
    "await fetch(\"/camera/select\", {\n"
    "    method: \"POST\",\n"
    "    body: JSON.stringify({ source: url })\n"
    "});\n"
    "status.textContent = `Stream conectado: ${url}`;"
)

heading("Problema encontrado", level=3)
body(
    "O OpenCV não suporta streams HLS (.m3u8) de forma confiável. Streams desse tipo nunca "
    "abriam, sem que o usuário recebesse qualquer erro. Além disso, todo o tráfego de vídeo "
    "passava desnecessariamente pelo servidor Python, aumentando latência e consumo de recursos "
    "mesmo quando o browser poderia tocar o stream diretamente."
)

heading("O que foi melhorado", level=3)
code_block(
    "const isHLS     = url.includes(\".m3u8\");\n"
    "const isYoutube = url.includes(\"youtube.com\") || url.includes(\"youtu.be\");\n\n"
    "if (isHLS) {\n"
    "    // hls.js toca direto no browser, sem passar pelo Python\n"
    "    switchToHLS(url);\n"
    "} else if (isYoutube) {\n"
    "    // Backend resolve a URL; browser toca o resultado com hls.js\n"
    "    const { resolved_url } = await fetch(\"/stream/resolve\", { body: ... });\n"
    "    switchToHLS(resolved_url);\n"
    "} else {\n"
    "    // RTSP / MJPEG: proxy pelo backend via OpenCV\n"
    "    await fetch(\"/camera/select\", { body: JSON.stringify({ source: url }) });\n"
    "}"
)

heading("Por que a nova versão é melhor", level=3)
body(
    "Cada tipo de stream é tratado pela tecnologia mais adequada: HLS é reproduzido diretamente "
    "no browser usando hls.js, que possui suporte nativo a adaptive bitrate e reconexão automática; "
    "YouTube é resolvido pelo backend (yt-dlp) e o resultado tocado pelo browser; RTSP e MJPEG "
    "continuam sendo processados pelo OpenCV no servidor. Isso reduz a carga no servidor, elimina "
    "a latência de proxy e suporta mais formatos de stream."
)

# ══════════════════════════════════════════════════════════════
# PARTE 4
# ══════════════════════════════════════════════════════════════
doc.add_page_break()
heading("Parte 4 — Implementação da Camada de Web Scraping")

heading("4.1 Justificativa e relevância", level=2)
body(
    "O AgroVision AI detecta pessoas, veículos e movimentações em propriedades rurais. Isoladamente, "
    "uma detecção registra apenas 'o que' foi visto e 'quando'. A camada de scraping adiciona contexto "
    "externo que permite ao operador — e ao assistente de IA — interpretar melhor os eventos:"
)
bullet("Uma detecção de pessoas em horário incomum pode estar relacionada a uma colheita emergencial motivada por previsão de chuva")
bullet("Movimentação intensa de veículos pode coincidir com notícias de alta no preço de commodities")
bullet("Alertas climáticos podem explicar comportamentos atípicos no campo")
body(
    "Os dois dados escolhidos — previsão do tempo e notícias do agronegócio — são os mais "
    "diretamente relevantes para o contexto rural e estão disponíveis publicamente e gratuitamente."
)

heading("4.2 Fontes utilizadas", level=2)
add_table(
    ["Dado", "Fonte", "Técnica", "API Key necessária?"],
    [
        ["Previsão do tempo (atual + 3 dias)", "Open-Meteo (open-meteo.com)", "API REST pública", "Não"],
        ["Notícias do agronegócio", "noticiasagricolas.com.br", "Web scraping com BeautifulSoup", "Não"],
    ]
)

doc.add_paragraph()
heading("4.3 Implementação — scraper.py", level=2)
body(
    "A camada de scraping foi implementada no arquivo scraper.py, completamente separado do app.py. "
    "O módulo expõe duas funções públicas e gerencia internamente cache, rate limiting e tratamento de erro."
)

heading("Estrutura do módulo", level=3)
code_block(
    "scraper.py\n"
    "  ├── _get_cache(key)        ← verifica cache por TTL\n"
    "  ├── _set_cache(key, data)  ← armazena resultado em cache\n"
    "  ├── _rate_ok(key)          ← rate limiting por endpoint\n"
    "  ├── fetch_weather(lat, lon) ← coleta clima via Open-Meteo\n"
    "  └── fetch_agro_news(limit)  ← scraping de notícias"
)

heading("Boas práticas implementadas", level=3)

body("Cache por TTL (10 minutos):", bold=True)
code_block(
    "CACHE_TTL = 600  # segundos\n\n"
    "def _get_cache(key: str):\n"
    "    entry = _CACHE.get(key)\n"
    "    if entry and (time.time() - entry[\"ts\"]) < CACHE_TTL:\n"
    "        return entry[\"data\"]\n"
    "    return None"
)

body("Rate limiting entre requisições reais:", bold=True)
code_block(
    "MIN_INTERVAL = 60  # segundos\n\n"
    "def _rate_ok(key: str) -> bool:\n"
    "    now = time.time()\n"
    "    if now - _LAST_REQUEST.get(key, 0) < MIN_INTERVAL:\n"
    "        return False\n"
    "    _LAST_REQUEST[key] = now\n"
    "    return True"
)

body("Timeout e tratamento de erro:", bold=True)
code_block(
    "REQUEST_TIMEOUT = 10  # segundos\n\n"
    "try:\n"
    "    resp = requests.get(url, timeout=REQUEST_TIMEOUT)\n"
    "    resp.raise_for_status()\n"
    "    # ... processa dados ...\n"
    "except Exception as exc:\n"
    "    return {\"error\": f\"Clima indisponível: {exc}\"}"
)

body("Scraping robusto com múltiplos seletores CSS:", bold=True)
code_block(
    "# Tenta diferentes seletores para resistir a mudanças de layout\n"
    "candidates = (\n"
    "    soup.select(\"div.item-noticia a\")\n"
    "    or soup.select(\"ul.list-noticias li a\")\n"
    "    or soup.select(\"article a\")\n"
    "    or soup.select(\"h2 a, h3 a\")\n"
    ")"
)

heading("Formato de saída estruturado (JSON)", level=3)
body("Clima (fetch_weather):")
code_block(
    "{\n"
    "  \"current\": {\n"
    "    \"temperature\": 19.4,\n"
    "    \"humidity\": 58,\n"
    "    \"precipitation\": 0.0,\n"
    "    \"windspeed\": 6.0\n"
    "  },\n"
    "  \"forecast\": [\n"
    "    {\"date\": \"2026-05-28\", \"max\": 19.4, \"min\": 11.4, \"rain\": 0.0},\n"
    "    {\"date\": \"2026-05-29\", \"max\": 19.5, \"min\": 13.1, \"rain\": 0.3}\n"
    "  ],\n"
    "  \"source\": \"open-meteo.com\"\n"
    "}"
)
body("Notícias (fetch_agro_news):")
code_block(
    "[\n"
    "  {\"title\": \"Soja: safra recorde impacta preços...\",\n"
    "   \"link\": \"https://www.noticiasagricolas.com.br/...\",\n"
    "   \"source\": \"noticiasagricolas.com.br\"},\n"
    "  ...\n"
    "]"
)

heading("4.4 Integração com o sistema", level=2)
body("A camada de scraping é integrada em três pontos:")
bullet("Endpoints da API: GET /scraping/weather e GET /scraping/news no app.py")
bullet("Dashboard: seção 'Clima Atual' e 'Notícias do Agronegócio' exibidas automaticamente ao carregar a página")
bullet("Assistente Ollama: o contexto do último evento é enviado ao modelo junto com os dados climáticos, permitindo respostas contextualizadas")

add_table(
    ["Endpoint", "Método", "Parâmetros", "Descrição"],
    [
        ["/scraping/weather", "GET", "lat, lon (opcional)", "Clima atual e previsão para 3 dias"],
        ["/scraping/news", "GET", "limit (padrão: 5)", "Últimas notícias do agronegócio"],
    ]
)

# ══════════════════════════════════════════════════════════════
# CONCLUSÃO
# ══════════════════════════════════════════════════════════════
doc.add_page_break()
heading("Conclusão")
body(
    "Esta revisão evidenciou que o código inicial, gerado com auxílio de IA, funcionava corretamente "
    "do ponto de vista funcional, mas apresentava problemas estruturais relevantes: ausência de "
    "separação de responsabilidades, vulnerabilidades de segurança e código bloqueante em contextos "
    "assíncronos."
)
doc.add_paragraph()
body(
    "As melhorias realizadas — timeout no yt-dlp, placeholder no stream MJPEG e roteamento "
    "inteligente de streams no cliente — demonstram que revisar código gerado por IA é etapa "
    "obrigatória antes de qualquer implantação. A IA acelera a prototipagem, mas não substitui "
    "o raciocínio do desenvolvedor sobre arquitetura, segurança e escalabilidade."
)
doc.add_paragraph()
body(
    "A camada de web scraping implementada em scraper.py é o exemplo mais claro dessa postura: "
    "ao invés de adicionar as funções diretamente no app.py (como a IA provavelmente sugeriria), "
    "optou-se por um módulo isolado com responsabilidade única, respeitando princípios de "
    "coesão e baixo acoplamento."
)

# ── Salva ──────────────────────────────────────────────────────
output_path = "Revisao_Arquitetura_AgroVision.docx"
doc.save(output_path)
print(f"Documento gerado: {output_path}")
